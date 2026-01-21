"""
Document Ingestion

Parses clinical documents from various formats:
- PDF (with OCR for scanned documents)
- DOCX
- Plain text
- HL7 messages

All processing is local - no external API calls.
"""

import io
import re
from pathlib import Path
from typing import BinaryIO

import structlog
from pydantic import BaseModel, Field

from app.core.security import generate_secure_id
from app.config import settings

logger = structlog.get_logger(__name__)


class DocumentMetadata(BaseModel):
    """Metadata about a parsed document."""
    id: str = Field(default_factory=generate_secure_id)
    filename: str
    file_type: str
    file_size: int
    
    # Extracted metadata
    title: str | None = None
    author: str | None = None
    created_date: str | None = None
    
    # Processing info
    pages: int = 1
    was_ocr_processed: bool = False
    processing_time_ms: int = 0
    
    # Content stats
    character_count: int = 0
    word_count: int = 0


class ParsedDocument(BaseModel):
    """Result of document parsing."""
    metadata: DocumentMetadata
    content: str
    
    # Structured sections if identifiable
    sections: dict[str, str] = {}
    
    # Extracted entities (if pre-processed)
    patient_info: dict | None = None
    dates: list[str] = []
    providers: list[str] = []


class DocumentParser:
    """
    Multi-format document parser.
    
    All processing happens locally for HIPAA compliance.
    """
    
    def __init__(self):
        self._tesseract_available = False
        self._check_tesseract()
    
    def _check_tesseract(self) -> None:
        """Check if Tesseract OCR is available."""
        try:
            import pytesseract
            if settings.tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = settings.tesseract_path
            # Try a simple test
            pytesseract.get_tesseract_version()
            self._tesseract_available = True
            logger.info("Tesseract OCR available")
        except Exception as e:
            logger.warning("Tesseract OCR not available", error=str(e))
    
    def parse(self, file: BinaryIO, filename: str) -> ParsedDocument:
        """
        Parse a document from a file-like object.
        
        Automatically detects format based on extension.
        """
        import time
        start_time = time.time()
        
        # Determine file type
        extension = Path(filename).suffix.lower()
        file_size = file.seek(0, 2)  # Seek to end
        file.seek(0)  # Reset
        
        content = ""
        was_ocr = False
        pages = 1
        
        if extension == ".pdf":
            content, pages, was_ocr = self._parse_pdf(file)
        elif extension in [".docx", ".doc"]:
            content = self._parse_docx(file)
        elif extension == ".txt":
            content = self._parse_text(file)
        elif extension == ".hl7":
            content = self._parse_hl7(file)
        else:
            # Try to parse as text
            try:
                content = file.read().decode("utf-8", errors="ignore")
            except Exception:
                raise ValueError(f"Unsupported file format: {extension}")
        
        processing_time = int((time.time() - start_time) * 1000)
        
        # Build metadata
        metadata = DocumentMetadata(
            filename=filename,
            file_type=extension,
            file_size=file_size,
            pages=pages,
            was_ocr_processed=was_ocr,
            processing_time_ms=processing_time,
            character_count=len(content),
            word_count=len(content.split()),
        )
        
        # Extract sections
        sections = self._identify_sections(content)
        
        # Extract dates
        dates = self._extract_dates(content)
        
        return ParsedDocument(
            metadata=metadata,
            content=content,
            sections=sections,
            dates=dates,
        )
    
    def _parse_pdf(self, file: BinaryIO) -> tuple[str, int, bool]:
        """
        Parse a PDF file.
        
        Uses text extraction first, falls back to OCR for scanned PDFs.
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file)
            pages = len(reader.pages)
            
            # Try text extraction first
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            content = "\n\n".join(text_parts)
            
            # If very little text extracted, try OCR
            if len(content.strip()) < 100 and self._tesseract_available:
                file.seek(0)
                content, _ = self._ocr_pdf(file)
                return content, pages, True
            
            return content, pages, False
            
        except Exception as e:
            logger.error("PDF parsing failed", error=str(e))
            raise ValueError(f"Failed to parse PDF: {e}")
    
    def _ocr_pdf(self, file: BinaryIO) -> tuple[str, int]:
        """OCR a scanned PDF using Tesseract."""
        try:
            import pytesseract
            from PIL import Image
            from pdf2image import convert_from_bytes
            
            # Convert PDF to images
            images = convert_from_bytes(file.read())
            
            text_parts = []
            for image in images:
                text = pytesseract.image_to_string(image)
                text_parts.append(text)
            
            return "\n\n".join(text_parts), len(images)
            
        except ImportError:
            logger.warning("pdf2image not available for OCR")
            return "", 0
        except Exception as e:
            logger.error("OCR failed", error=str(e))
            return "", 0
    
    def _parse_docx(self, file: BinaryIO) -> str:
        """Parse a DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file)
            paragraphs = [p.text for p in doc.paragraphs]
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            logger.error("DOCX parsing failed", error=str(e))
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    def _parse_text(self, file: BinaryIO) -> str:
        """Parse a plain text file."""
        try:
            content = file.read()
            # Try different encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return content.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error("Text parsing failed", error=str(e))
            raise ValueError(f"Failed to parse text file: {e}")
    
    def _parse_hl7(self, file: BinaryIO) -> str:
        """
        Parse an HL7 v2.x message.
        
        Extracts human-readable content from segments.
        """
        try:
            import hl7
            
            content = file.read().decode("utf-8", errors="ignore")
            message = hl7.parse(content)
            
            # Extract key segments
            readable_parts = []
            
            # Patient info from PID
            if "PID" in message:
                pid = message["PID"][0]
                readable_parts.append(f"Patient: {pid[5]}")
                readable_parts.append(f"DOB: {pid[7]}")
            
            # Diagnosis from DG1
            for dg1 in message.segments("DG1"):
                readable_parts.append(f"Diagnosis: {dg1[3]} - {dg1[4]}")
            
            # Notes from NTE
            for nte in message.segments("NTE"):
                if len(nte) > 3:
                    readable_parts.append(f"Note: {nte[3]}")
            
            # OBX observations
            for obx in message.segments("OBX"):
                if len(obx) > 5:
                    readable_parts.append(f"Observation: {obx[3]} = {obx[5]}")
            
            return "\n".join(readable_parts)
            
        except Exception as e:
            logger.error("HL7 parsing failed", error=str(e))
            # Return raw content
            try:
                return file.read().decode("utf-8", errors="ignore")
            except:
                return ""
    
    def _identify_sections(self, content: str) -> dict[str, str]:
        """
        Identify common clinical document sections.
        
        Looks for headers like:
        - CHIEF COMPLAINT
        - HISTORY OF PRESENT ILLNESS
        - ASSESSMENT
        - PLAN
        etc.
        """
        sections = {}
        
        # Common section headers
        section_patterns = [
            r"(?i)CHIEF\s+COMPLAINT[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)HISTORY\s+OF\s+PRESENT\s+ILLNESS[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)HPI[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)PAST\s+MEDICAL\s+HISTORY[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)PMH[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)PHYSICAL\s+EXAM(?:INATION)?[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)ASSESSMENT[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)PLAN[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)DIAGNOS(?:IS|ES)[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)PROCEDURE[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
            r"(?i)OPERATIVE\s+REPORT[:\s]*(.*?)(?=\n[A-Z]{2,}|\Z)",
        ]
        
        section_names = [
            "chief_complaint",
            "history_present_illness",
            "history_present_illness",
            "past_medical_history",
            "past_medical_history",
            "physical_exam",
            "assessment",
            "plan",
            "diagnoses",
            "procedure",
            "operative_report",
        ]
        
        for pattern, name in zip(section_patterns, section_names):
            match = re.search(pattern, content, re.DOTALL)
            if match:
                section_text = match.group(1).strip()
                if section_text and name not in sections:
                    sections[name] = section_text[:2000]  # Limit size
        
        return sections
    
    def _extract_dates(self, content: str) -> list[str]:
        """Extract dates from document content."""
        date_patterns = [
            r"\d{1,2}/\d{1,2}/\d{2,4}",  # MM/DD/YYYY
            r"\d{4}-\d{2}-\d{2}",  # YYYY-MM-DD
            r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}",
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            dates.extend(matches)
        
        # Deduplicate and limit
        return list(set(dates))[:20]


# Singleton parser instance
_document_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    """Get the document parser instance."""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser

